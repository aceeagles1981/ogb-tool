"""
doc_extract.py — Document extraction for OGB Tool
Reads attachments, classifies doc_type + doc_stage, extracts structured terms.
Stores extracted data only — original files are never persisted.
Uses app.py's anthropic_text() and extract_json_object() for AI calls.
"""

import io
import json
import re
import logging
from datetime import datetime

logger = logging.getLogger("og_backend")

# ── Text extraction by file type ──────────────────────────────────────────────

MAX_RAW_TEXT = 100_000  # ~100KB cap on stored raw text


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> tuple:
    """
    Extract text from file bytes based on filename extension.
    Returns (text, was_truncated).
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    text = ''
    truncated = False

    try:
        if ext == 'pdf':
            text = _extract_pdf(file_bytes)
        elif ext in ('docx', 'doc'):
            text = _extract_docx(file_bytes)
        elif ext in ('xlsx', 'xls'):
            text = _extract_xlsx(file_bytes)
        elif ext in ('txt', 'csv', 'tsv'):
            text = file_bytes.decode('utf-8', errors='replace')
        elif ext in ('msg',):
            text = _extract_msg_text(file_bytes)
        elif ext in ('eml',):
            text = _extract_eml_text(file_bytes)
        elif ext in ('htm', 'html'):
            text = _strip_html(file_bytes.decode('utf-8', errors='replace'))
        else:
            try:
                text = file_bytes.decode('utf-8', errors='strict')
            except UnicodeDecodeError:
                text = f"[Binary file: {filename}, {len(file_bytes)} bytes, cannot extract text]"

        if len(text) > MAX_RAW_TEXT:
            text = text[:MAX_RAW_TEXT]
            truncated = True

    except Exception as e:
        logger.error(f"Text extraction failed for {filename}: {e}")
        text = f"[Extraction failed for {filename}: {str(e)}]"

    return text.strip(), truncated


def _extract_pdf(data: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ''
            tables = page.extract_tables() or []
            table_text = ''
            for table in tables:
                for row in table:
                    if row:
                        table_text += ' | '.join(str(cell or '') for cell in row) + '\n'
            if page_text:
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")
            if table_text:
                text_parts.append(f"--- Page {i+1} Tables ---\n{table_text}")
    return '\n\n'.join(text_parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else '' for cell in row]
            if any(cells):
                parts.append(' | '.join(cells))
    wb.close()
    return '\n'.join(parts)


def _extract_msg_text(data: bytes) -> str:
    try:
        import extract_msg
        import tempfile, os
        with tempfile.NamedTemporaryFile(delete=False, suffix=".msg") as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            msg = extract_msg.Message(tmp_path)
            parts = []
            if msg.subject: parts.append(f"Subject: {msg.subject}")
            if msg.sender: parts.append(f"From: {msg.sender}")
            if msg.date: parts.append(f"Date: {msg.date}")
            if msg.body: parts.append(f"\n{msg.body}")
            return '\n'.join(parts)
        finally:
            try: os.unlink(tmp_path)
            except OSError: pass
    except Exception as e:
        return f"[MSG extraction failed: {e}]"


def _extract_eml_text(data: bytes) -> str:
    import email
    from email import policy as ep
    msg = email.message_from_bytes(data, policy=ep.default)
    parts = []
    parts.append(f"Subject: {msg.get('subject', '')}")
    parts.append(f"From: {msg.get('from', '')}")
    parts.append(f"Date: {msg.get('date', '')}")
    body = msg.get_body(preferencelist=('plain', 'html'))
    if body:
        content = body.get_content()
        if body.get_content_type() == 'text/html':
            content = _strip_html(content)
        parts.append(f"\n{content}")
    return '\n'.join(parts)


def _strip_html(html: str) -> str:
    text = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


# ── Attachment extraction from email files ────────────────────────────────────

def extract_attachments_from_msg(file_bytes: bytes) -> list:
    attachments = []
    try:
        import extract_msg, tempfile, os
        with tempfile.NamedTemporaryFile(delete=False, suffix=".msg") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            msg = extract_msg.Message(tmp_path)
            for att in (msg.attachments or []):
                data = getattr(att, 'data', None)
                name = getattr(att, 'longFilename', None) or getattr(att, 'filename', None)
                if data and name:
                    attachments.append({'filename': str(name), 'data': data, 'size': len(data)})
        finally:
            try: os.unlink(tmp_path)
            except OSError: pass
    except Exception as e:
        logger.error(f"Failed to extract .msg attachments: {e}")
    return attachments


def extract_attachments_from_eml(file_bytes: bytes) -> list:
    import email
    from email import policy as ep
    attachments = []
    try:
        msg = email.message_from_bytes(file_bytes, policy=ep.default)
        for part in msg.walk():
            cd = part.get('Content-Disposition', '')
            if 'attachment' in cd or (part.get_filename() and part.get_content_maintype() != 'multipart'):
                filename = part.get_filename()
                data = part.get_payload(decode=True)
                if filename and data:
                    attachments.append({'filename': filename, 'data': data, 'size': len(data)})
    except Exception as e:
        logger.error(f"Failed to extract .eml attachments: {e}")
    return attachments


# ── AI Classification & Extraction Prompts ────────────────────────────────────

CLASSIFY_SYSTEM = """You are a Lloyd's of London marine cargo insurance document classifier.
Classify the document into EXACTLY one doc_type and one doc_stage.

doc_type options: slip, endorsement, debit_note, credit_note, inspection_report, sov,
proposal_form, wording, correspondence, certificate, claim_advice, survey_report,
loss_record, account_statement, unknown

doc_stage options:
- indicated: terms under discussion (producer proposed, broker draft, underwriter quote)
- firm: terms agreed (may have outstanding admin items)
- endorsement: amendment to firm/bound terms

source_party: who created this document (e.g. producer name, "OGB", underwriter name)

Respond ONLY with JSON:
{"doc_type": "...", "doc_stage": "...", "source_party": "...", "confidence": 0.85, "reasoning": "..."}
"""

EXTRACT_TERMS_SYSTEM = """You are a Lloyd's marine cargo insurance terms extractor. You are expert at reading reinsurance slips, MRCs, and proposal forms — including those formatted as Word tables.

CRITICAL RULES:
- Extract EVERY number you find. Do not return null or omit a field if the document contains the data.
- If a field says "to be informed" or is blank, set it to null — do NOT set it to 0 or NaN.
- Rates must be decimals: 0.20% = 0.0020. Never return a percentage as 0.20 or 20.
- If no premium or rate is stated in the document, set premium to null. NEVER invent a premium.
- If the document says "No claims" or "Nil" for loss history, return "loss_history": "Nil claims reported" — this IS data, not missing data.
- Multiple limits (ocean, ground, air, storage, NatCat) should ALL be captured in the limits object.
- Multiple deductibles (transit, stock, NatCat) should ALL be captured.
- Look for data in TABLES — Word documents often have key:value pairs in table cells.
- The cedant/reinsured is the original insurance company, NOT the insured or the broker.

Return ONLY valid JSON:
{
  "insured": "full legal name",
  "goods": "full goods description as stated",
  "product": "Marine Cargo | STP | WHLL | FFL | Project Cargo | Other",
  "cedant": "original insurer / reinsured company name",
  "annual_turnover": {"amount": number_or_null, "currency": "USD"},
  "limits": {
    "any_one_conveyance": number_or_null,
    "any_one_location": number_or_null,
    "any_one_loss": number_or_null,
    "aggregate": number_or_null,
    "ocean": number_or_null,
    "ground": number_or_null,
    "air": number_or_null,
    "storage": number_or_null,
    "natcat": number_or_null,
    "unnamed_location": number_or_null,
    "containers_on_deck": number_or_null,
    "currency": "USD"
  },
  "deductible": {
    "transit": number_or_null,
    "stock": number_or_null,
    "natcat": "description e.g. 10% coinsurance",
    "currency": "USD",
    "basis": "each and every loss"
  },
  "rate": {"transit": decimal_or_null, "stock": decimal_or_null, "flat": decimal_or_null},
  "premium": {"amount": number_or_null, "currency": "USD", "basis": "minimum_and_deposit | adjustable | flat | null"},
  "locations": [
    {"name": "full address", "country": "...", "max_value": number, "average_value": number_or_null}
  ],
  "transits": {
    "imports_from": ["country or region"],
    "exports_to": ["country: city"],
    "domestic": "description or null"
  },
  "incoterms": {"stated": true_or_false, "terms": {"CIF": percent, "FOB": percent}},
  "valuation_basis": "description e.g. C&F +10% imports, Book value +10% storage",
  "perils": "ICC(A) | ICC(B) | ICC(C) | All Risks | Named Perils",
  "war_cover": true_or_false,
  "strikes_cover": true_or_false,
  "exclusions": ["exact exclusion clause names/descriptions"],
  "conditions": ["exact condition descriptions"],
  "warranties": ["exact warranty descriptions"],
  "subjectivities": [{"item": "...", "status": "outstanding | received | waived"}],
  "period": {"inception": "YYYY-MM-DD or null", "expiry": "YYYY-MM-DD or null", "duration_months": number_or_null},
  "brokerage": number_or_null,
  "loss_history": "Nil claims reported | description of claims | null if genuinely not provided",
  "payment_terms": "description of premium payment terms",
  "law_jurisdiction": "governing law and jurisdiction",
  "approved_adjusters": ["names"],
  "placement": {"layers": [{"name": "Primary", "limit": number, "excess": number,
    "lines": [{"market": "...", "pct": number, "role": "lead | follow"}]}]},
  "endorsement_ref": "EN1 if applicable",
  "endorsement_type": "AP | RP | extension | loss_payee | correction",
  "endorsement_changes": "description",
  "loss_payees": ["names"],
  "policy_ref": "UMR or policy reference",
  "notes": "any other important terms not captured above"
}
"""

EXTRACT_SURVEY_SYSTEM = """You are a marine cargo insurance survey report analyst. You may receive reports in any language — extract data regardless of language.

Extract ALL specific findings. Do not summarise vaguely — name exact deficiencies and recommendations.
For fire protection, security, and storage: give the specific equipment present or absent.

Return ONLY valid JSON:
{
  "surveyor": "inspector name and company",
  "survey_date": "YYYY-MM-DD",
  "locations_inspected": ["full address of each location"],
  "overall_rating": "satisfactory | needs_improvement | unsatisfactory",
  "fire_protection": {
    "rating": "good | adequate | poor",
    "sprinklers": true_or_false,
    "extinguishers": true_or_false,
    "fire_hose": true_or_false,
    "fire_doors": true_or_false,
    "fire_alarm": true_or_false,
    "emergency_lighting": true_or_false,
    "notes": "specific details about what is present or missing"
  },
  "security": {
    "rating": "good | adequate | poor",
    "cctv": true_or_false,
    "alarm_system": true_or_false,
    "security_guard": true_or_false,
    "perimeter_fence": true_or_false,
    "notes": "specific details"
  },
  "storage_conditions": {
    "rating": "good | adequate | poor",
    "notes": "specific details about racking, stacking, housekeeping"
  },
  "construction": {
    "walls": "material description",
    "roof": "material description",
    "floor": "material description",
    "doors": "material description",
    "year_built": number_or_null,
    "floors": number_or_null,
    "notes": "any structural concerns"
  },
  "natural_hazard_exposure": {
    "earthquake": true_or_false,
    "flood": true_or_false,
    "windstorm": true_or_false,
    "notes": "drainage issues, proximity to water, seismic zone"
  },
  "recommendations": ["each specific recommendation as stated in the report"],
  "deficiencies": ["each specific deficiency found — be exact, e.g. 'No sprinkler system installed', 'No emergency lighting', 'Fire protection rated malo (bad)'"],
  "occupancy": "description of building usage and number of employees",
  "pml_estimate": {"amount": number_or_null, "currency": "USD", "scenario": "description"},
  "max_values_at_location": number_or_null,
  "notes": "any other key observations including partial occupancy, shared buildings, uncontrolled adjacent risks"
}
"""


def _ai_call(system: str, user_text: str, api_key: str, model: str, max_tokens: int = 2000) -> dict:
    """Call AI and parse JSON response. Uses app.py helpers if available, falls back to direct."""
    try:
        from app import anthropic_text, extract_json_object
        raw = anthropic_text({
            'model': model,
            'max_tokens': max_tokens,
            'system': system,
            'messages': [{'role': 'user', 'content': user_text}]
        })
        return extract_json_object(raw)
    except ImportError:
        # Fallback: direct call (shouldn't happen in production)
        import requests as req
        resp = req.post(
            'https://api.anthropic.com/v1/messages',
            headers={'Content-Type': 'application/json', 'x-api-key': api_key, 'anthropic-version': '2023-06-01'},
            json={'model': model, 'max_tokens': max_tokens, 'system': system,
                  'messages': [{'role': 'user', 'content': user_text}]},
            timeout=60
        )
        resp.raise_for_status()
        data = resp.json()
        text = ''.join(b.get('text', '') for b in data.get('content', []) if b.get('type') == 'text').strip()
        text = re.sub(r'^```json\s*', '', text)
        text = re.sub(r'\s*```$', '', text)
        return json.loads(text)


# ── Main extraction pipeline ─────────────────────────────────────────────────

def process_attachment(file_bytes, filename, email_context, api_key, model='claude-haiku-4-5-20251001'):
    """Full pipeline for one attachment: extract text → classify → extract terms/survey."""
    result = {
        'filename': filename,
        'file_type': filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'unknown',
        'raw_text': '', 'raw_text_truncated': False,
        'classification': None, 'terms': None, 'survey_findings': None, 'error': None
    }

    # Step 1: Extract text
    try:
        raw_text, truncated = extract_text_from_bytes(file_bytes, filename)
        result['raw_text'] = raw_text
        result['raw_text_truncated'] = truncated
    except Exception as e:
        result['error'] = f"Text extraction failed: {str(e)}"
        return result

    if not raw_text or raw_text.startswith('[Binary file:') or raw_text.startswith('[Extraction failed'):
        result['error'] = f"No readable text in {filename}"
        return result

    # Step 2: Classify
    classify_prompt = f"EMAIL CONTEXT:\n{email_context[:2000]}\n\nDOCUMENT TEXT:\n{raw_text[:8000]}"
    try:
        classification = _ai_call(CLASSIFY_SYSTEM, classify_prompt, api_key, model, 500)
        result['classification'] = classification
    except Exception as e:
        logger.error(f"Classification failed for {filename}: {e}")
        result['classification'] = {
            'doc_type': 'unknown', 'doc_stage': 'indicated',
            'source_party': 'unknown', 'confidence': 0.0,
            'reasoning': f'Classification failed: {str(e)}'
        }

    doc_type = result['classification'].get('doc_type', 'unknown')

    # Step 3: Extract terms or survey findings
    extract_prompt = f"Document classified as: {doc_type} / {result['classification'].get('doc_stage', 'indicated')}\n\nEMAIL CONTEXT:\n{email_context[:2000]}\n\nDOCUMENT TEXT:\n{raw_text[:15000]}"
    try:
        if doc_type in ('inspection_report', 'survey_report'):
            result['survey_findings'] = _ai_call(EXTRACT_SURVEY_SYSTEM, extract_prompt, api_key, model, 2000)
        elif doc_type not in ('correspondence', 'unknown'):
            result['terms'] = _ai_call(EXTRACT_TERMS_SYSTEM, extract_prompt, api_key, model, 3000)
    except Exception as e:
        logger.error(f"Terms extraction failed for {filename}: {e}")
        result['error'] = f"Terms extraction partial failure: {str(e)}"

    return result


def process_all_attachments(email_file_bytes, email_filename, email_body_text, api_key, model='claude-haiku-4-5-20251001'):
    """Extract and process all attachments from an email file."""
    ext = email_filename.rsplit('.', 1)[-1].lower() if '.' in email_filename else ''

    if ext == 'msg':
        attachments = extract_attachments_from_msg(email_file_bytes)
    elif ext == 'eml':
        attachments = extract_attachments_from_eml(email_file_bytes)
    else:
        return [process_attachment(email_file_bytes, email_filename, email_body_text, api_key, model)]

    results = []
    skip_exts = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'ico', 'svg'}
    for att in attachments:
        if att['size'] < 500:
            continue
        att_ext = att['filename'].rsplit('.', 1)[-1].lower() if '.' in att['filename'] else ''
        if att_ext in skip_exts:
            continue
        results.append(process_attachment(att['data'], att['filename'], email_body_text, api_key, model))

    return results
